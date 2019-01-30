'''
This is example code for Python Package project by KB for CS550 HO. This code shows some of the most important and useful functions of that python-docx has to offer. In order to run this program with the image, simply replace the image name from one on your computer or download the image from the link in the IBook. 

On my honor, I have neither given nor recieved unauthorized aid. 

1/30/19 

'''

from docx import Document
from docx.shared import Inches

document = Document()

styles = document.styles


document.add_heading('Welcome to Python-Docx', 0)

p = document.add_paragraph('Python-Docx can create text with many dfferent properties such as: ')
p.add_run('bold').bold = True
p.add_run(' and ')
p.add_run('italic.').italic = True


document.add_heading('It can also create heading levels', level=1)

document.add_paragraph('And sytles such as intense quote'
	, style='Intense Quote')

document.add_paragraph(
    'Dont forget about lists!', style='List Bullet').underline = True 



document.add_paragraph(
    'or ordered lists!', style='List Number'
)

document.add_paragraph("and of course pictures:")

document.add_picture('opengraph-icon-200x200.png', width=Inches(1.25))

paragraph = document.add_paragraph("And Tables")

table = document.add_table(1, 4)
cell1 = table.cell(0, 1)
cell1.text = 'item1'
cell2 = table.cell(0,2)
cell2.text = "item2"





document.save('example.docx')         